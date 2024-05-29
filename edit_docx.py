from docx import Document
import doc2pdf
import numpy as np

def read_edit_docx(value_arr):
    
    fullName = ""
    address = ""

    document = Document('template.docx')

    table = document.tables[0]

    table.cell(1,1).text = value_arr["lst_name"].upper()
    table.cell(1,4).text = value_arr["fst_name"].upper()
    table.cell(2,1).text = value_arr["fa_fst_name"].upper()
    table.cell(2,4).text = value_arr["grandfa_fst_name"].upper()
    table.cell(3,1).text = value_arr["ma_fst_name"].upper()
    table.cell(4,1).text = value_arr["sex"]
    table.cell(4,4).text = value_arr["id_number"]
    table.cell(5,1).text = value_arr["nationality"]
    table.cell(5,4).text = value_arr["religion"]
    table.cell(6,1).text = value_arr["date_nationality"]
    table.cell(6,4).text = value_arr["date_religion"]
    table.cell(7,1).text = value_arr["marital"]
    table.cell(7,4).text = value_arr["date_marital"]
    table.cell(8,1).text = value_arr["stateBirth"]
    table.cell(8,4).text = value_arr["city_birth"]
    table.cell(9,1).text = value_arr["hebrew_birth"]
    table.cell(9,4).text = value_arr["gregorian_birth"]
    table.cell(10,1).text = value_arr["date_entrance"]
    table.cell(10,4).text = value_arr["status"]
    table.cell(11,1).text = value_arr["registration_date"]
    table.cell(12,1).text = value_arr["address"]
    table.cell(13,1).text = value_arr["date_entry"]

    fullName = table.cell(1, 1).text + "  " + table.cell(1, 4).text
    address = table.cell(12, 1).text

    for i in np.arange(1,11):
        if(table.cell(i, 1).text == ""):
            table.cell(i, 0).text = ""

        if(table.cell(i, 4).text == ""):
            table.cell(i, 3).text = ""
    
    if(table.cell(11, 1).text == ""):
            table.cell(11, 0).text = ""
    if(table.cell(12, 1).text == ""):
            table.cell(12, 0).text = ""
    if(table.cell(13, 1).text == ""):
            table.cell(13, 0).text = ""


    table = document.tables[1]
    table.cell(0,2).text = value_arr["issuedOn"]

    # for paragraph in document.paragraphs:
    #     if("Notes:" in paragraph.text):
    #         index = paragraph.text.index("Notes:")
    #         paragraph.text = paragraph.text[:index + len("Notes:")] + "   " + value_arr["notes"] 

    for paragraph in document.paragraphs:
        if("Authority in" in paragraph.text):
            index = paragraph.text.index("Authority in")
            paragraph.text = paragraph.text[:index + len("Authority in")] + " " + value_arr["authorityIn"] 
    
    last_paragraph = document.paragraphs[-1]
    last_paragraph.text = last_paragraph.text + "\n" + fullName + "\n" + "\n" + address



    document.save('static/result_document.docx')
    doc2pdf.convert_docx2pdf()